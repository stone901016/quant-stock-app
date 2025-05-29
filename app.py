import os
import tempfile
import pandas as pd
from flask import Flask, request, render_template_string, send_file
from docx import Document
from FinMind.data import DataLoader
from datetime import datetime, timedelta

app = Flask(__name__)
TOKEN = "eyJ0eXAiOiJKV1QiLCJhbGciOiJIUzI1NiJ9.eyJkYXRlIjoiMjAyNS0wNS0yOCAwMToyODoxMyIsInVzZXJfaWQiOiJqYW1lczkwMTAxNiIsImlwIjoiMTE4LjE1MC42My45OSJ9.QWxBrJYWM_GNDpTyvAyR2frCPwB4e7HP_Kj_KEX2tVs"

@app.route('/')
def index():
    with open("templates/index.html", encoding="utf-8") as f:
        return render_template_string(f.read())

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        pe_ratio = float(request.form.get("pe_ratio", 15))
        pb_ratio = float(request.form.get("pb_ratio", 1.5))
        eps_growth = float(request.form.get("eps_growth", 10))
        roe = float(request.form.get("roe", 15))
        debt_ratio = float(request.form.get("debt_ratio", 50))
        price_3m = float(request.form.get("price_3m", 0.1))
        std_1y = float(request.form.get("std_1y", 0.2))
    except:
        return "❌ 請輸入正確的數字格式。"

    api = DataLoader()
    api.login_by_token(api_token=TOKEN)

    try:
        stock_list = api.taiwan_stock_info()
        stock_list = stock_list[stock_list["stock_id"].isin([
            "2330", "2454", "2317", "2308", "2382", "2891", "2881", "2882", "2303", "2412",
            "2886", "3711", "2884", "1216", "1301", "2002", "2880", "1326", "5871", "2207",
            "2883", "4938", "2912", "3008", "3034", "3037", "2603", "1101", "6415", "6669",
            "1590", "3481", "3045", "1402", "2885", "9910", "2357", "2609", "9904", "2345",
            "2379", "5876", "2301", "2892", "2395", "2408", "6414", "4958", "2801", "9914"
        ])]
    except Exception as e:
        return f"❌ 抓取股票清單失敗：{e}"

    result_list = []
    today = datetime.today()
    start = (today - timedelta(days=365)).strftime("%Y-%m-%d")
    end = today.strftime("%Y-%m-%d")

    print(f"📊 共 {len(stock_list)} 檔股票待分析")

    for _, row in stock_list.iterrows():
        stock_id = row["stock_id"]
        stock_name = row["stock_name"]

        try:
            print(f"🔍 正在分析 {stock_id} - {stock_name}")
            price_df = api.taiwan_stock_daily(stock_id=stock_id, start_date=start, end_date=end)
            if price_df.empty or len(price_df) < 65:
                print("❌ 價格資料不足")
                continue

            price_df["return"] = price_df["close"].pct_change()
            std = price_df["return"].std() * (252 ** 0.5)
            pct_3m = (price_df.iloc[-1]["close"] / price_df.iloc[-63]["close"]) - 1

            fin_df = api.taiwan_stock_financial_statement(
                stock_id=stock_id,
                start_date="2022-01-01",
                end_date="2024-12-31"
            )

            if fin_df.empty or "EPS" not in fin_df.columns:
                print("❌ 找不到財報資料或缺少欄位")
                continue

            fin_df = fin_df.sort_values("date", ascending=False)
            recent_4q = fin_df.head(4)

            eps_base = recent_4q["EPS"].sum()
            roe = recent_4q["ROE"].mean()
            debt_ratio_val = recent_4q["DebtRatio"].mean()
            pe_val = recent_4q["PER"].mean()
            pb_val = recent_4q["PBR"].mean()

            eps_target = eps_growth / 100 * eps_base

            print(f"📉 近3月報酬率: {pct_3m:.2%}, 年化波動: {std:.2%}")
            print(f"📊 財報資料: PER={pe_val}, PBR={pb_val}, EPS={eps_base}, ROE={roe}, 負債比={debt_ratio_val}")

            if (
                pb_val < pb_ratio and
                pe_val < pe_ratio and
                eps_base > 0 and eps_base > eps_target and
                roe > roe and
                debt_ratio_val < debt_ratio and
                pct_3m > price_3m and
                std < std_1y
            ):
                print(f"✅ {stock_id} 通過條件")
                result_list.append({
                    "symbol": stock_id,
                    "name": stock_name,
                    "pe_ratio": pe_val,
                    "pb_ratio": pb_val,
                    "eps": eps_base,
                    "roe": roe,
                    "debt_ratio": debt_ratio_val,
                    "price_3m": round(pct_3m, 2),
                    "std_1y": round(std, 4)
                })

        except Exception as e:
            print(f"❌ {stock_id} 錯誤：{e}")
            continue

    df = pd.DataFrame(result_list)
    doc = Document()
    doc.add_heading("多因子選股分析報告", 0)

    if df.empty:
        doc.add_paragraph("沒有符合條件的個股。")
    else:
        for _, row in df.iterrows():
            line = (
                f"{row['symbol']} - {row['name']} | "
                f"PE: {row['pe_ratio']} PB: {row['pb_ratio']} EPS: {row['eps']} "
                f"ROE: {row['roe']} 負債比: {row['debt_ratio']} "
                f"3月漲: {row['price_3m']} 波動: {row['std_1y']}"
            )
            doc.add_paragraph(line)

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)

    return send_file(temp.name, as_attachment=True, download_name="選股分析報告.docx")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
